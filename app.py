from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel


BASE_DIR = Path(__file__).parent
TEMPLATE_PATH = BASE_DIR / "Data grid template.docx"
OUTPUT_DIR = BASE_DIR / "generated"
OUTPUT_DIR.mkdir(exist_ok=True)

ALLOWED_LETTERS = tuple("ABCDEFGHIJKLMNOPQR")
FIELD_ORDER = [
    "date",
    "photograph",
    "client",
    "elevation",
    "grid_letter",
    "grid_number",
    "probe",
    "substrate",
    "comments",
]


class DataRow(BaseModel):
    grid_letter: str = ""
    grid_number: str = ""
    probe: str = ""
    substrate: str = ""
    comments: str = ""


class GenerateRequest(BaseModel):
    document_name: str
    date: str = ""
    photograph: str = ""
    client: str = ""
    elevation: str = ""
    rows: list[DataRow] = []


app = FastAPI(title="EIFS Voice Grid Tool")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.mount("/generated", StaticFiles(directory=str(OUTPUT_DIR)), name="generated")


def _sanitize_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9 _.-]", "", name).strip()
    if not cleaned:
        cleaned = "Data grid output"
    return cleaned[:120]


def _validate_grid_letter(value: str) -> bool:
    value = value.strip().upper()
    pattern = r"^(?:[A-R]|[A-R]-[A-R])$"
    return bool(re.match(pattern, value))


def _validate_grid_number(value: str) -> bool:
    value = value.strip()
    single = r"^(?:[1-9]|1[0-4])$"
    ranged = r"^(?:[1-9]|1[0-4])-(?:[1-9]|1[0-4])$"
    return bool(re.match(single, value) or re.match(ranged, value))


def _validate_probe(value: str) -> bool:
    value = value.strip()
    if not value.isdigit():
        return False
    number = int(value)
    return 10 <= number <= 40


def _validate_substrate(value: str) -> bool:
    return value.strip().upper() in {"F", "M", "S"}


def _set_cell_text(cell, value: str) -> None:
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _replace_label_in_runs(document: Document, label: str, value: str) -> bool:
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if label in run.text:
                run.text = f"{label} {value}" if value else label
                return True
    return False


def _find_grid_table(document: Document):
    for table in document.tables:
        if not table.rows:
            continue
        header = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if any("grid letter" in text for text in header) and any("grid no" in text for text in header):
            return table
    return None


def _resolve_column_map(table) -> dict[str, int]:
    header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
    mapping: dict[str, int] = {}

    for idx, text in enumerate(header_cells):
        if "grid letter" in text:
            mapping["grid_letter"] = idx
        elif "grid no" in text:
            mapping["grid_number"] = idx
        elif "probe" in text:
            mapping["probe"] = idx
        elif "substrate" in text:
            mapping["substrate"] = idx
        elif "comment" in text:
            mapping["comments"] = idx

    # Some templates have a blank substrate header in the 4th column.
    if "substrate" not in mapping and len(header_cells) >= 5:
        mapping["substrate"] = 3

    required = {"grid_letter", "grid_number", "probe", "substrate", "comments"}
    missing = required - set(mapping)
    if missing:
        raise ValueError(f"Template columns are missing: {', '.join(sorted(missing))}")

    return mapping


def _append_row_clone(table):
    new_row = deepcopy(table.rows[-1]._tr)
    table._tbl.append(new_row)


def _validate_rows(rows: list[DataRow]) -> None:
    for i, row in enumerate(rows, start=1):
        if row.grid_letter and not _validate_grid_letter(row.grid_letter):
            raise ValueError(f"Row {i}: grid letter must be A-R or A-R format with one dash (example A-C).")
        if row.grid_number and not _validate_grid_number(row.grid_number):
            raise ValueError(f"Row {i}: grid number must be 1-14 or range like 1-14.")
        if row.probe and not _validate_probe(row.probe):
            raise ValueError(f"Row {i}: probe must be a number between 10 and 40.")
        if row.substrate and not _validate_substrate(row.substrate):
            raise ValueError(f"Row {i}: substrate must be F, M, or S.")


def _populate_document(payload: GenerateRequest, destination: Path) -> None:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")

    document = Document(str(TEMPLATE_PATH))

    _replace_label_in_runs(document, "Date:", payload.date)
    _replace_label_in_runs(document, "Photograph:", payload.photograph)
    _replace_label_in_runs(document, "Client:", payload.client)
    _replace_label_in_runs(document, "Elevation:", payload.elevation)

    grid_table = _find_grid_table(document)
    if grid_table is None:
        raise ValueError("Could not find grid data table in template.")

    column_map = _resolve_column_map(grid_table)
    rows = payload.rows or [DataRow()]
    _validate_rows(rows)

    required_table_rows = len(rows) + 1
    while len(grid_table.rows) < required_table_rows:
        _append_row_clone(grid_table)

    for idx, row in enumerate(rows, start=1):
        target = grid_table.rows[idx]
        _set_cell_text(target.cells[column_map["grid_letter"]], row.grid_letter.strip().upper())
        _set_cell_text(target.cells[column_map["grid_number"]], row.grid_number.strip())
        _set_cell_text(target.cells[column_map["probe"]], row.probe.strip())
        _set_cell_text(target.cells[column_map["substrate"]], row.substrate.strip().upper())
        _set_cell_text(target.cells[column_map["comments"]], row.comments.strip())

    document.save(str(destination))


@app.get("/api/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/api/generate")
def generate_document(payload: GenerateRequest):
    safe_name = _sanitize_filename(payload.document_name)
    output_path = OUTPUT_DIR / f"{safe_name}.docx"

    try:
        _populate_document(payload, output_path)
    except (ValueError, FileNotFoundError) as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:  # pragma: no cover - defensive guard
        raise HTTPException(status_code=500, detail=f"Unexpected error: {exc}") from exc

    return FileResponse(
        str(output_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=output_path.name,
    )


@app.get("/")
def root():
    return FileResponse(str(BASE_DIR / "index.html"))
